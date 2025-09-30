from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Type
from dataclasses import dataclass
from pathlib import Path
import mimetypes
from loguru import logger
from PIL import Image
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
import pandas
import ezdxf
import time
import os
import io


@dataclass
class ParserResult:
    """Результат парсинга документа"""
    success: bool
    text: str
    error_message: str
    metadata: Dict[str, Any]
    file_path: str
    
    def is_success(self) -> bool:
        return self.success
    
    def get_text(self) -> str:
        return self.text
    
    def get_error(self) -> str:
        return self.error_message
    
    def get_metadata(self) -> Dict:
        return self.metadata


class BaseParser(ABC):
    """Базовый интерфейс парсера"""
    
    @abstractmethod
    def parse(self, file_path: str, **params) -> ParserResult:
        """Парсинг файла"""
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """Получить поддерживаемые расширения"""
        pass
    
    def validate_file(self, file_path: str) -> bool:
        """Валидация файла"""
        if not Path(file_path).exists():
            return False
        
        ext = Path(file_path).suffix.lower()
        return ext in self.get_supported_extensions()


class FileValidator:
    """Валидатор файлов"""
    
    @staticmethod
    def validate_file_exists(file_path: str) -> bool:
        return Path(file_path).exists()
    
    @staticmethod
    def validate_file_size(file_path: str, max_size: int) -> bool:
        if not Path(file_path).exists():
            return False
        return Path(file_path).stat().st_size <= max_size
    
    @staticmethod
    def validate_file_type(file_path: str, expected_extensions: List[str]) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in expected_extensions
    
    @staticmethod
    def get_file_mime_type(file_path: str) -> str:
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type or 'application/octet-stream'


class PDFParser(BaseParser):
    """Парсер PDF документов"""
    
    def __init__(self, use_ocr: bool = True, extract_images: bool = False, 
                 ocr_language: str = "rus+eng"):
        self.use_ocr = use_ocr
        self.extract_images = extract_images
        self.ocr_language = ocr_language
    
    def parse(self, file_path: str, **params) -> ParserResult:

        """ Парсинг PDF файла """

        start_time = time.time()
        
        try:
            # Извлекаем параметры
            use_ocr = params.get('use_ocr', self.use_ocr)
            ocr_language = params.get('ocr_language', self.ocr_language)
            pages = params.get('pages', None)
            
            # Проверяем файл
            if not os.path.exists(file_path):
                return ParserResult(
                    success=False,
                    text="",
                    error_message=f"Файл не найден: {file_path}",
                    metadata={},
                    file_path=file_path
                )
            
            # Выбираем метод парсинга
            if use_ocr:
                import pytesseract
                from PIL import Image
                import pymupdf as fitz
                text, metadata = self._extract_with_ocr(file_path, pages, ocr_language)
                method = "ocr"
            else:
                text, metadata = self._extract_with_pymupdf(file_path, pages)
                method = "text_extraction"
            
            # Извлекаем метаданные
            pdf_metadata = self._extract_pdf_metadata(file_path)
            
            # Формируем итоговые метаданные
            final_metadata = {
                'parser': 'PDFParser',
                'method': method,
                'use_ocr': use_ocr,
                'processing_time_sec': round(time.time() - start_time, 2),
                'text_length': len(text),
                **metadata,
                **pdf_metadata
            }
            
            return ParserResult(
                success=True,
                text=text,
                error_message="",
                metadata=final_metadata,
                file_path=file_path
            )
            
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return ParserResult(
                success=False,
                text="",
                error_message=str(e),
                metadata={},
                file_path=file_path
            )
    
    def _extract_with_pymupdf(self, file_path: str, pages: List[int] = None) -> tuple:
        """Извлечение текста с помощью PyMuPDF (fitz) - САМЫЙ БЫСТРЫЙ"""
        import pymupdf as fitz
        
        text = ""
        metadata = {}
        
        with fitz.open(file_path) as doc:
            metadata['total_pages'] = len(doc)
            metadata['is_encrypted'] = doc.is_encrypted
            
            # Определяем страницы для обработки
            if pages:
                page_indices = [p-1 for p in pages if 1 <= p <= len(doc)]
            else:
                page_indices = range(len(doc))
            
            for page_num in page_indices:
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text += f"--- Страница {page_num + 1} ---\n{page_text}\n"
            
            metadata['pages_processed'] = len(page_indices)
        
        return text, metadata
    
    def _extract_with_ocr(self, file_path: str, pages: List[int] = None, language: str = "rus+eng") -> tuple:
        """Извлечение текста с помощью OCR"""
        import pymupdf as fitz
        import pytesseract
        from PIL import Image
        
        text = ""
        metadata = {
            'ocr_language': language,
            'ocr_engine': 'tesseract'
        }
        
        with fitz.open(file_path) as doc:
            metadata['total_pages'] = len(doc)
            
            # Определяем страницы для обработки
            if pages:
                page_indices = [p-1 for p in pages if 1 <= p <= len(doc)]
            else:
                page_indices = range(len(doc))
            
            ocr_pages_count = 0
            
            for page_num in page_indices:
                page = doc[page_num]
                
                # Сначала пробуем извлечь обычный текст
                page_text = page.get_text()
                if page_text.strip():
                    text += f"--- Страница {page_num + 1} (ТЕКСТ) ---\n{page_text}\n"
                else:
                    # Если текста нет - используем OCR
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Увеличиваем разрешение
                    img_data = pix.tobytes("png")
                    
                    with Image.open(io.BytesIO(img_data)) as img:
                        ocr_text = pytesseract.image_to_string(img, lang=language)
                        if ocr_text.strip():
                            text += f"--- Страница {page_num + 1} (OCR) ---\n{ocr_text}\n"
                            ocr_pages_count += 1
            
            metadata['pages_processed'] = len(page_indices)
            metadata['ocr_pages'] = ocr_pages_count
        
        return text, metadata
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Извлечение метаданных PDF"""
        import pymupdf as fitz
        
        metadata = {}
        
        try:
            with fitz.open(file_path) as doc:
                pdf_metadata = doc.metadata
                metadata.update({
                    'author': pdf_metadata.get('author', ''),
                    'title': pdf_metadata.get('title', ''),
                    'subject': pdf_metadata.get('subject', ''),
                    'keywords': pdf_metadata.get('keywords', ''),
                    'creator': pdf_metadata.get('creator', ''),
                    'producer': pdf_metadata.get('producer', ''),
                    'creation_date': pdf_metadata.get('creationDate', ''),
                    'modification_date': pdf_metadata.get('modDate', ''),
                })
        except:
            pass  # Игнорируем ошибки метаданных
        
        # Добавляем информацию о файле
        file_stats = os.stat(file_path)
        metadata.update({
            'file_size_bytes': file_stats.st_size,
        })
        
        return metadata
    
    def get_supported_extensions(self) -> List[str]:
        return ['.pdf']


class DOCParser(BaseParser):
    """Парсер для .doc файлов с правильной обработкой"""
    
    def parse(self, file_path: str, **params) -> ParserResult:
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext != '.doc':
                return ParserResult(
                    success=False,
                    text="",
                    error_message="DOCParser supports only .doc files",
                    metadata={},
                    file_path=file_path
                )
            
            # Пробуем методы в порядке приоритета
            return self._try_all_doc_methods(file_path)
            
        except Exception as e:
            logger.error(f"DOC parsing error: {e}")
            return ParserResult(
                success=False,
                text="",
                error_message=str(e),
                metadata={},
                file_path=file_path
            )
    
    def _try_all_doc_methods(self, file_path: str) -> ParserResult:
        """Пробуем все доступные методы парсинга .doc"""
        
        methods = [
            self._parse_with_olefile,      # Метод 1: OLE структура
            self._parse_with_antiword,     # Метод 2: antiword
            self._parse_with_strings,      # Метод 3: strings command
            self._parse_smart_binary,      # Метод 4: Умное бинарное чтение
        ]
        
        for method in methods:
            result = method(file_path)
            if result.success:
                logger.info(f"DOC parsed successfully with {result.metadata.get('method')}")
                return result
        
        return ParserResult(
            success=False,
            text="",
            error_message="All .doc parsing methods failed. File may be corrupted, encrypted, or in unsupported format.",
            metadata={},
            file_path=file_path
        )
    
    def _parse_with_olefile(self, file_path: str) -> ParserResult:
        """Парсинг .doc через OLE структуру (самый надежный для старых Word)"""
        try:
            import olefile
            
            if not olefile.isOleFile(file_path):
                return ParserResult(success=False, text="", error_message="Not a valid OLE file", metadata={}, file_path=file_path)
            
            ole = olefile.OleFileIO(file_path)
            text_parts = []
            
            # Пытаемся извлечь текст из различных потоков
            streams = ['WordDocument', 'Table', 'Data', 'SummaryInformation']
            
            for stream in streams:
                if ole.exists(stream):
                    try:
                        data = ole.openstream(stream).read()
                        # Пытаемся извлечь текст из данных
                        stream_text = self._extract_text_from_bytes(data)
                        if stream_text:
                            text_parts.append(stream_text)
                    except:
                        continue
            
            ole.close()
            
            if text_parts:
                text = '\n'.join(text_parts)
                return ParserResult(
                    success=True,
                    text=text,
                    error_message="",
                    metadata={'method': 'olefile', 'parser': 'DOCParser'},
                    file_path=file_path
                )
            else:
                return ParserResult(success=False, text="", error_message="No text found in OLE streams", metadata={}, file_path=file_path)
                
        except ImportError:
            return ParserResult(success=False, text="", error_message="olefile not installed", metadata={}, file_path=file_path)
        except Exception as e:
            return ParserResult(success=False, text="", error_message=str(e), metadata={}, file_path=file_path)
    
    def _parse_with_antiword(self, file_path: str) -> ParserResult:
        """Парсинг через antiword (требует установки antiword)"""
        try:
            import subprocess
            
            # Проверяем доступность antiword
            result = subprocess.run(['antiword', '-v'], capture_output=True, text=True)
            if result.returncode != 0:
                return ParserResult(success=False, text="", error_message="antiword not available", metadata={}, file_path=file_path)
            
            # Запускаем antiword
            result = subprocess.run(
                ['antiword', '-m', 'UTF-8.txt', file_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore',
                timeout=30
            )
            
            if result.returncode == 0 and result.stdout.strip():
                return ParserResult(
                    success=True,
                    text=result.stdout,
                    error_message="",
                    metadata={'method': 'antiword', 'parser': 'DOCParser'},
                    file_path=file_path
                )
            else:
                return ParserResult(success=False, text="", error_message="antiword failed", metadata={}, file_path=file_path)
                
        except Exception as e:
            return ParserResult(success=False, text="", error_message=str(e), metadata={}, file_path=file_path)
    
    def _parse_with_strings(self, file_path: str) -> ParserResult:
        """Использование strings команды для извлечения текста"""
        try:
            import subprocess
            
            # Запускаем strings для извлечения текстовых последовательностей
            result = subprocess.run(
                ['strings', '-n', '10', file_path],  # Минимум 10 символов
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore',
                timeout=30
            )
            
            if result.returncode == 0 and result.stdout.strip():
                lines = [line.strip() for line in result.stdout.split('\n') if len(line.strip()) > 10]
                
                # Фильтруем только строки с русскими буквами
                russian_lines = []
                for line in lines:
                    russian_count = sum(1 for c in line if 'а' <= c <= 'я' or 'А' <= c <= 'Я')
                    if russian_count > len(line) * 0.2:  # Хотя бы 20% русских букв
                        russian_lines.append(line)
                
                if russian_lines:
                    text = '\n'.join(russian_lines)
                    return ParserResult(
                        success=True,
                        text=text,
                        error_message="",
                        metadata={'method': 'strings', 'parser': 'DOCParser'},
                        file_path=file_path
                    )
            
            return ParserResult(success=False, text="", error_message="No readable text found with strings", metadata={}, file_path=file_path)
            
        except Exception as e:
            return ParserResult(success=False, text="", error_message=str(e), metadata={}, file_path=file_path)
    
    def _parse_smart_binary(self, file_path: str) -> ParserResult:
        """Умное чтение бинарного файла с поиском текстовых блоков"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Ищем текстовые блоки в бинарных данных
            text_blocks = self._find_text_blocks(content)
            
            if text_blocks:
                text = '\n'.join(text_blocks)
                return ParserResult(
                    success=True,
                    text=text,
                    error_message="",
                    metadata={'method': 'smart_binary', 'parser': 'DOCParser', 'blocks_found': len(text_blocks)},
                    file_path=file_path
                )
            else:
                return ParserResult(success=False, text="", error_message="No text blocks found", metadata={}, file_path=file_path)
                
        except Exception as e:
            return ParserResult(success=False, text="", error_message=str(e), metadata={}, file_path=file_path)
    
    def _find_text_blocks(self, content: bytes) -> List[str]:
        """Поиск текстовых блоков в бинарных данных"""
        text_blocks = []
        
        # Кодировки для попытки
        encodings = ['utf-16-le', 'utf-16-be', 'cp1251', 'cp866', 'koi8-r']
        
        for encoding in encodings:
            try:
                # Пытаемся декодировать весь файл
                decoded = content.decode(encoding, errors='ignore')
                
                # Разбиваем на строки и фильтруем
                lines = []
                for line in decoded.split('\n'):
                    line = line.strip()
                    if 10 <= len(line) <= 1000:  # Разумная длина
                        # Проверяем на текст (буквы, цифры, пунктуация)
                        text_chars = sum(1 for c in line if c.isalnum() or c in ' .,!?;-')
                        if text_chars > len(line) * 0.6:  # Хотя бы 60% текстовых символов
                            lines.append(line)
                
                if lines:
                    text_blocks.extend(lines)
                    break
                    
            except Exception:
                continue
        
        return text_blocks
    
    def _extract_text_from_bytes(self, data: bytes) -> str:
        """Извлечение текста из байтовых данных"""
        text = ""
        
        for encoding in ['utf-16-le', 'utf-8', 'cp1251', 'cp866']:
            try:
                decoded = data.decode(encoding, errors='ignore')
                # Ищем последовательности с буквами
                lines = [line.strip() for line in decoded.split('\n') if len(line.strip()) > 5 and any(c.isalpha() for c in line)]
                if lines:
                    text = '\n'.join(lines)
                    break
            except:
                continue
        
        return text
    
    def get_supported_extensions(self) -> List[str]:
        return ['.doc']


class DOCXParser(BaseParser):
    """Парсер ТОЛЬКО для DOCX файлов"""
    
    def parse(self, file_path: str, **params) -> ParserResult:
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext != '.docx':
                return ParserResult(
                    success=False,
                    text="",
                    error_message=f"DOCXParser supports only .docx files. Use DOCParser for .doc files.",
                    metadata={},
                    file_path=file_path
                )
            
            from docx import Document
            doc = Document(file_path)
            text = self._parse_document_structure(doc)
            
            metadata = {
                'parser': 'DOCXParser',
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables)
            }
            
            return ParserResult(
                success=True,
                text=text,
                error_message="",
                metadata=metadata,
                file_path=file_path
            )
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return ParserResult(
                success=False,
                text="",
                error_message=str(e),
                metadata={},
                file_path=file_path
            )
    
    def _parse_document_structure(self, doc) -> str:
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        return text
    
    def get_supported_extensions(self) -> List[str]:
        return ['.docx']


class XLSXParser(BaseParser):
    """Парсер Excel документов"""
    
    def __init__(self, read_formulas: bool = False, sheet_names: Optional[List[str]] = None):
        self.read_formulas = read_formulas
        self.sheet_names = sheet_names
    
    def parse(self, file_path: str, **params) -> ParserResult:
        try:
            import pandas as pd
            
            if self.sheet_names:
                excel_file = pd.read_excel(file_path, sheet_name=self.sheet_names)
            else:
                excel_file = pd.read_excel(file_path, sheet_name=None)
            
            text = self._parse_sheets(excel_file)
            
            metadata = {
                'parser': 'XLSXParser',
                'sheets_count': len(excel_file) if isinstance(excel_file, dict) else 1
            }
            
            return ParserResult(
                success=True,
                text=text,
                error_message="",
                metadata=metadata,
                file_path=file_path
            )
        except Exception as e:
            logger.error(f"XLSX parsing error: {e}")
            return ParserResult(
                success=False,
                text="",
                error_message=str(e),
                metadata={},
                file_path=file_path
            )
    
    def _parse_sheets(self, excel_file) -> str:
        """Парсинг листов Excel"""
        text = ""
        
        if isinstance(excel_file, dict):
            for sheet_name, df in excel_file.items():
                text += f"\n--- Лист: {sheet_name} ---\n"
                text += df.to_string() + "\n"
        else:
            text = excel_file.to_string()
        
        return text
    
    def get_supported_extensions(self) -> List[str]:
        return ['.xlsx', '.xls']


class PlainTextParser(BaseParser):
    """Парсер текстовых файлов"""
    
    def __init__(self, encodings: Optional[List[str]] = None):
        self.encodings = encodings or ['utf-8', 'cp1251', 'latin-1']
    
    def parse(self, file_path: str, **params) -> ParserResult:
        try:
            text = self._try_encodings(file_path)
            
            metadata = {
                'parser': 'PlainTextParser',
                'encoding': 'detected'
            }
            
            return ParserResult(
                success=True,
                text=text,
                error_message="",
                metadata=metadata,
                file_path=file_path
            )
        except Exception as e:
            logger.error(f"Text parsing error: {e}")
            return ParserResult(
                success=False,
                text="",
                error_message=str(e),
                metadata={},
                file_path=file_path
            )
    
    def _try_encodings(self, file_path: str) -> str:
        """Попытка прочитать с разными кодировками"""
        for encoding in self.encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # Последняя попытка с игнорированием ошибок
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def get_supported_extensions(self) -> List[str]:
        return ['.txt', '.log', '.csv', '.md']

class DWGParser(BaseParser):
    pass


class DXFParser(BaseParser):
    """Парсер DXF файлов (AutoCAD Drawing Exchange Format)"""
    
    def __init__(self, extract_metadata: bool = True, extract_blocks: bool = True,
                 search_in_blocks: bool = True, search_in_layouts: bool = True):
        self.extract_metadata = extract_metadata
        self.extract_blocks = extract_blocks
        self.search_in_blocks = search_in_blocks
        self.search_in_layouts = search_in_layouts
    
    def parse(self, file_path: str, **params) -> ParserResult:

        """ Парсинг DXF файла """
        
        try:
            # Извлекаем параметры поиска
            search_in_blocks = params.get('search_in_blocks', self.search_in_blocks)
            search_in_layouts = params.get('search_in_layouts', self.search_in_layouts)
            search_in_entities = params.get('search_in_entities', True)
            deep_search = params.get('deep_search', True)
            
            # Открываем DXF файл
            doc = ezdxf.readfile(file_path)
            
            # Ищем текст в РАЗНЫХ местах
            text_content = ""
            
            if search_in_entities:
                text_content += self._extract_from_entities(doc)
            
            if search_in_blocks:
                text_content += self._extract_from_blocks(doc, deep_search)
            
            if search_in_layouts:
                text_content += self._extract_from_layouts(doc)
            
            # Извлекаем метаданные
            metadata = self._extract_metadata(doc) if self.extract_metadata else {}
            
            # Формируем итоговый текст
            final_text = self._format_output(text_content, metadata, params)
            
            return ParserResult(
                success=True,
                text=final_text,
                error_message="",
                metadata=metadata,
                file_path=file_path
            )
            
        except Exception as e:
            return ParserResult(
                success=False,
                text="",
                error_message=f"DXF parsing error: {str(e)}",
                metadata={},
                file_path=file_path
            )
    
    def _extract_from_entities(self, doc) -> str:
        """Извлечение текста из entities (прямых объектов)"""
        text_parts = []
        msp = doc.modelspace()
        
        # Счетчики найденных объектов
        counters = {'TEXT': 0, 'MTEXT': 0, 'ATTDEF': 0, 'ATTRIB': 0}
        
        # TEXT entities
        for text in msp.query('TEXT'):
            if text.dxf.text and text.dxf.text.strip():
                text_parts.append(f"TEXT: {text.dxf.text}")
                counters['TEXT'] += 1
        
        # MTEXT entities (многострочный текст)
        for mtext in msp.query('MTEXT'):
            if mtext.text and mtext.text.strip():
                text_parts.append(f"MTEXT: {mtext.text}")
                counters['MTEXT'] += 1
        
        # ATTDEF (определения атрибутов)
        for attdef in msp.query('ATTDEF'):
            if attdef.dxf.tag and attdef.dxf.default_value:
                text_parts.append(f"ATTR_DEF: {attdef.dxf.tag} = {attdef.dxf.default_value}")
                counters['ATTDEF'] += 1
        
        # ATTRIB (атрибуты вставленных блоков)
        for attrib in msp.query('ATTRIB'):
            if attrib.dxf.text and attrib.dxf.text.strip():
                text_parts.append(f"ATTR: {attrib.dxf.text}")
                counters['ATTRIB'] += 1
        
        # Добавляем статистику
        if any(counters.values()):
            text_parts.append(f"\n📊 Найдено в entities: TEXT={counters['TEXT']}, MTEXT={counters['MTEXT']}, ATTDEF={counters['ATTDEF']}, ATTRIB={counters['ATTRIB']}")
        
        return "\n".join(text_parts) + "\n" if text_parts else ""
    
    def _extract_from_blocks(self, doc, deep_search: bool = True) -> str:
        """Извлечение текста из блоков"""
        text_parts = []
        total_blocks_searched = 0
        total_text_found = 0
        
        for block in doc.blocks:
            # Пропускаем системные блоки
            if block.name.startswith('*'):
                continue
                
            total_blocks_searched += 1
            block_text_parts = []
            
            # Ищем текст внутри блока
            for entity in block:
                entity_text = self._extract_text_from_entity(entity)
                if entity_text:
                    block_text_parts.append(f"  - {entity_text}")
                    total_text_found += 1
            
            # Если в блоке нашли текст - добавляем в результат
            if block_text_parts:
                text_parts.append(f"🔷 БЛОК: {block.name}")
                text_parts.extend(block_text_parts)
                text_parts.append("")  # Пустая строка для разделения
        
        # Добавляем статистику по блокам
        if total_blocks_searched > 0:
            text_parts.append(f"📊 Поиск в блоках: проверено {total_blocks_searched} блоков, найдено текста в {total_text_found} местах")
        
        return "\n".join(text_parts) + "\n" if text_parts else ""
    
    def _extract_from_layouts(self, doc) -> str:
        """Извлечение текста из layout'ов (Paper Space)"""
        text_parts = []
        
        for layout in doc.layouts:
            # Пропускаем Model Space (уже обработали)
            if layout.name == 'Model':
                continue
                
            layout_text_parts = []
            
            # Ищем текст в layout'е
            for entity in layout:
                entity_text = self._extract_text_from_entity(entity)
                if entity_text:
                    layout_text_parts.append(f"  - {entity_text}")
            
            # Если в layout'е нашли текст
            if layout_text_parts:
                text_parts.append(f"📄 LAYOUT: {layout.name}")
                text_parts.extend(layout_text_parts)
                text_parts.append("")
        
        return "\n".join(text_parts) + "\n" if text_parts else ""
    
    def _extract_text_from_entity(self, entity) -> str:
        """Извлечение текста из конкретного entity"""
        try:
            if entity.dxftype() == 'TEXT' and entity.dxf.text and entity.dxf.text.strip():
                return f"TEXT: {entity.dxf.text}"
            elif entity.dxftype() == 'MTEXT' and entity.text and entity.text.strip():
                return f"MTEXT: {entity.text}"
            elif entity.dxftype() == 'ATTDEF' and entity.dxf.tag and entity.dxf.default_value:
                return f"ATTR_DEF: {entity.dxf.tag} = {entity.dxf.default_value}"
            elif entity.dxftype() == 'ATTRIB' and entity.dxf.text and entity.dxf.text.strip():
                return f"ATTR: {entity.dxf.text}"
        except:
            pass
        return ""
    
    def _extract_metadata(self, doc) -> Dict[str, Any]:
        """Извлечение метаданных DXF файла"""
        metadata = {
            'dxf_version': str(doc.dxfversion),
            'layers_count': len(doc.layers),
            'blocks_count': len(doc.blocks),
            'entities_count': len(doc.modelspace()),
            'layouts_count': len(doc.layouts) - 1,  # -1 потому что Model тоже layout
            'file_units': str(doc.header.get('$INSUNITS', 'Unknown')),
        }
        
        # Статистика по типам объектов
        msp = doc.modelspace()
        entity_stats = {
            'TEXT': len(msp.query('TEXT')),
            'MTEXT': len(msp.query('MTEXT')),
            'ATTDEF': len(msp.query('ATTDEF')),
            'ATTRIB': len(msp.query('ATTRIB')),
            'INSERT': len(msp.query('INSERT')),  # Вставки блоков
        }
        metadata['entity_statistics'] = entity_stats
        
        # Информация о слоях
        layers_info = []
        for layer in doc.layers:
            layers_info.append({
                'name': layer.dxf.name,
                'color': layer.dxf.color,
                'is_off': layer.is_off(),
            })
        metadata['layers'] = layers_info
        
        # Информация о блоках
        blocks_info = []
        for block in doc.blocks:
            if not block.name.startswith('*'):
                blocks_info.append({
                    'name': block.name,
                    'entities_count': len(block),
                })
        metadata['blocks'] = blocks_info
        
        return metadata
    
    def _format_output(self, text_content: str, metadata: Dict, params: Dict) -> str:
        """Форматирование итогового текста"""
        output_parts = []
        
        # Добавляем метаданные
        if params.get('include_metadata', True):
            output_parts.append("=== МЕТАДАННЫЕ DXF ===")
            output_parts.append(f"Версия DXF: {metadata.get('dxf_version', 'Unknown')}")
            output_parts.append(f"Количество слоев: {metadata.get('layers_count', 0)}")
            output_parts.append(f"Количество блоков: {metadata.get('blocks_count', 0)}")
            output_parts.append(f"Количество объектов в Model: {metadata.get('entities_count', 0)}")
            
            # Статистика объектов
            stats = metadata.get('entity_statistics', {})
            output_parts.append(f"📊 Объекты: TEXT={stats.get('TEXT', 0)}, MTEXT={stats.get('MTEXT', 0)}, ATTDEF={stats.get('ATTDEF', 0)}, ATTRIB={stats.get('ATTRIB', 0)}")
            output_parts.append("")
        
        # Добавляем текстовое содержимое
        if text_content.strip():
            output_parts.append("=== ТЕКСТОВОЕ СОДЕРЖИМОЕ ===")
            output_parts.append(text_content)
        else:
            output_parts.append("=== ТЕКСТ НЕ НАЙДЕН ===")
            output_parts.append("Текстовые объекты не найдены в:")
            output_parts.append("- Model Space entities")
            output_parts.append("- Блоках (blocks)")
            output_parts.append("- Layout'ах (Paper Space)")
            output_parts.append("")
            output_parts.append("💡 Возможные причины:")
            output_parts.append("- Текст находится во вставленных блоках (INSERT)")
            output_parts.append("- Файл содержит только геометрию без текста")
            output_parts.append("- Текст в специализированных объектах")
        
        return "\n".join(output_parts)
    
    def get_supported_extensions(self) -> List[str]:
        return ['.dxf']

class ImageOCRParser(BaseParser):
    """Парсер изображений с OCR"""
    
    def __init__(self, ocr_language: str = 'rus+eng+ita+spa', image_quality: str = 'high'):
        self.ocr_language = ocr_language
        self.image_quality = image_quality
    
    def parse(self, file_path: str, **params) -> ParserResult:
        
        try:
            image = Image.open(file_path)
            preprocessed = self._preprocess_image(image)
            text = self._perform_ocr(preprocessed)
            
            metadata = {
                'parser': 'ImageOCRParser',
                'language': self.ocr_language,
                'image_size': image.size
            }
            
            return ParserResult(
                success=True,
                text=text,
                error_message="",
                metadata=metadata,
                file_path=file_path
            )
        except Exception as e:
            logger.error(f"OCR parsing error: {e}")
            return ParserResult(
                success=False,
                text="",
                error_message=str(e),
                metadata={},
                file_path=file_path
            )
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Предобработка изображения для OCR"""
        # Конвертация в grayscale
        return image.convert('L')
    
    def _perform_ocr(self, image: Image.Image) -> str:
        """Выполнение OCR"""
        return pytesseract.image_to_string(image, lang=self.ocr_language)
    
    def get_supported_extensions(self) -> List[str]:
        return ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']


class DWGParser(BaseParser):
    """Парсер DWG/DXF файлов (чертежи)"""
    
    def __init__(self, extract_metadata: bool = True, extract_geometry: bool = False):
        self.extract_metadata = extract_metadata
        self.extract_geometry = extract_geometry
    
    def parse(self, file_path: str, **params) -> ParserResult:
        try:
            doc = ezdxf.readfile(file_path)
            text = self._extract_text_entities(doc)
            
            metadata = self._parse_dwg_metadata(doc) if self.extract_metadata else {}
            metadata['parser'] = 'DWGParser'
            
            return ParserResult(
                success=True,
                text=text,
                error_message="",
                metadata=metadata,
                file_path=file_path
            )
        except Exception as e:
            logger.error(f"DWG parsing error: {e}")
            return ParserResult(
                success=False,
                text="",
                error_message=str(e),
                metadata={},
                file_path=file_path
            )
    
    def _parse_dwg_metadata(self, doc) -> Dict:
        """Извлечение метаданных DWG"""
        return {
            'dxf_version': doc.dxfversion,
            'layers_count': len(doc.layers)
        }
    
    def _extract_text_entities(self, doc) -> str:
        """Извлечение текстовых сущностей"""
        text = ""
        msp = doc.modelspace()
        
        for entity in msp:
            if entity.dxftype() == 'TEXT':
                text += entity.dxf.text + "\n"
            elif entity.dxftype() == 'MTEXT':
                text += entity.text + "\n"
        
        return text
    
    def get_supported_extensions(self) -> List[str]:
        return ['.dwg', '.dxf']


class DepartmentConfig:
    """Конфигурация парсеров для отдела"""
    
    def __init__(self, department_name: str, max_file_size: int = 100 * 1024 * 1024, timeout: int = 300):
        self.department_name = department_name
        self.allowed_parsers: Dict[str, List[Type[BaseParser]]] = {}
        self.max_file_size = max_file_size
        self.timeout = timeout
        self._parser_priorities: Dict[str, Dict[Type[BaseParser], int]] = {}
    
    def get_allowed_extensions(self) -> List[str]:
        return list(self.allowed_parsers.keys())
    
    def is_parser_allowed(self, parser_class: Type[BaseParser]) -> bool:
        for parsers in self.allowed_parsers.values():
            if parser_class in parsers:
                return True
        return False
    
    def get_parser_priority(self, extension: str, parser_class: Type[BaseParser]) -> int:
        return self._parser_priorities.get(extension, {}).get(parser_class, 999)
    
    def register_parser(self, extension: str, parser_class: Type[BaseParser], priority: int = 10):
        if extension not in self.allowed_parsers:
            self.allowed_parsers[extension] = []
        
        if parser_class not in self.allowed_parsers[extension]:
            self.allowed_parsers[extension].append(parser_class)
        
        if extension not in self._parser_priorities:
            self._parser_priorities[extension] = {}
        self._parser_priorities[extension][parser_class] = priority


class ParserRegistry:

    """Реестр парсеров"""
    
    def __init__(self):
        self.global_registry: Dict[str, List[Type[BaseParser]]] = {}
        self.department_registries: Dict[str, DepartmentConfig] = {}
    

    def register_global_parser(self, extension: str, parser: Type[BaseParser]):
        if extension not in self.global_registry:
            self.global_registry[extension] = []
        if parser not in self.global_registry[extension]:
            self.global_registry[extension].append(parser)
    

    def register_department_parser(self, department: str, extension: str, parser: list[Type[BaseParser]]):
        if department not in self.department_registries:
            self.department_registries[department] = DepartmentConfig(department)
        
        self.department_registries[department].register_parser(extension, parser)
    

    def get_parsers_for_department(self, department: str, extension: str) -> List[Type[BaseParser]]:
        if department in self.department_registries:
            dept_config = self.department_registries[department]
            parsers = dept_config.allowed_parsers.get(extension, [])
            if parsers:
                return sorted(parsers, key=lambda p: dept_config.get_parser_priority(extension, p))
        
        return self.global_registry.get(extension, [])
    

    def get_all_supported_extensions(self) -> List[str]:
        extensions = set(self.global_registry.keys())
        for dept_config in self.department_registries.values():
            extensions.update(dept_config.get_allowed_extensions())
        return list(extensions)


class ParserManager:

    """Менеджер парсеров с поддержкой отделов"""
    
    def __init__(self):
        self.parser_registry = ParserRegistry()
        self.parser_instances: Dict[Type[BaseParser], BaseParser] = {}
        self.file_validator = FileValidator()
        
        self._init_default_parsers()
    

    def _init_default_parsers(self):

        """Инициализация стандартных парсеров"""

        default_parsers = {
            'pdf': PDFParser,
            'docx': DOCParser,
            'doc': DOCParser,
            'xlsx': XLSXParser,
            'xls': XLSXParser,
            'txt': PlainTextParser,
            'dxf': DXFParser,
            'dwg': DWGParser,
            'png': ImageOCRParser,
            'jpg': ImageOCRParser,
        }

        return default_parsers
    
            
    def _parser_extension(self, file_path: str) -> str:

        """Получение расширения файла"""

        return file_path.lower().split('.')[-1] if file_path else None
    

    def _find_parser_in_registry(self, extension: str) -> Optional[Type[BaseParser]]:

        """Поиск подходящего парсера в реестре"""

        return self._init_default_parsers().get(extension)
    

    def _save_parser_instance(self, parser_class: Type[BaseParser]):

        """Сохранение экземпляра парсера"""

        if parser_class not in self.parser_instances:
            self.parser_instances[parser_class] = parser_class()

    def _ransfer_selected_parser(self, file_path: str, parser_class: Type[BaseParser]):

        """Передача файла выбранному парсеру"""

        if parser_class is None:
            return None
        parser_instance = parser_class()
        return parser_instance.parse(file_path)


if __name__ == "__main__":
    path = "C:\\Users\\kulikovma\\Pictures\\Screenshots\\Снимок экрана 2025-09-30 162205.png"
    manager = ParserManager()
    ext = manager._parser_extension(path)
    print(ext)
    find = manager._find_parser_in_registry(ext)
    print(find)
    print(manager._ransfer_selected_parser(path, find))