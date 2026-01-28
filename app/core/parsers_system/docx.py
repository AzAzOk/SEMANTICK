from typing import List
from docx import Document
from .super_class import BaseParser, ParserResult
from loguru import logger

class DOCXParser(BaseParser):

    """Парсер для DOCX файлов"""
    
    def parse(self, file_path: str, **params) -> ParserResult:
        try:
            doc = Document(file_path)
            text = self._parse_document_structure(doc)
            
            metadata = {
                'parser': 'DOCXParser',
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables)
            }
            
            return ParserResult(
                success=True,
                text=text,
                error_message="",
                metadata=metadata,
                file_path=file_path
            )
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return ParserResult(
                success=False,
                text="",
                error_message=str(e),
                metadata={},
                file_path=file_path
            )
    

    def _parse_document_structure(self, doc) -> str:
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        return text
    

    def get_supported_extensions(self) -> List[str]:
        return ['.docx']